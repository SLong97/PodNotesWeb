import streamlit as st
import requests
from streamlit_player import st_player
import components.authenticate as authenticate
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import Frame, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT
import io
from docx import Document
import os
import boto3
from dotenv import load_dotenv
import re

st.set_page_config(page_title="Audio File Transcription", page_icon="📂")

# Check authentication
authenticate.set_st_state_vars()

# Add login/logout buttons
if st.session_state["authenticated"]:
    st.sidebar.markdown(" ")
    authenticate.button_logout()
else:
    st.sidebar.markdown(" ")
    authenticate.button_login()

st.write('<style> .css-1fcdlhc.e1s6o5jp0 {width: 900px !important;} </style>', unsafe_allow_html=True)

st.markdown("# Audio File Transcription")

st.write(
    """Upload your audio files and recieve precise transcriptions accompanied by concise summaries. 
    Facilitating the transformation of lectures, interviews, or any audio content into easily digestible text, making it more accessible and comprehensible."""
)

if(st.session_state["authenticated"]):

    load_dotenv()
    S3_BUCKET = os.getenv("S3_BUCKET")
    BUCKET_REGION = os.getenv("BUCKET_REGION")

    s3_client = boto3.client('s3', region_name=BUCKET_REGION)
    cognito_client = boto3.client('cognito-identity', region_name=os.getenv("COGNITO_REGION"))

    def get_cognito_identity(jwt_token):
        response = cognito_client.get_id(
            IdentityPoolId=os.getenv("COGNITO_IDENTITY_POOL_ID"),
            Logins={
                os.getenv("COGNITO_IDENTITY_PROVIDER"): jwt_token
            }
        )
        return response["IdentityId"]
    
    user_identity = get_cognito_identity(st.session_state["user_id"])

    st.write("")
    audio_file = st.file_uploader("Upload Audio", type=["wav", "mp3", "m4a"])

    PROCESS_UPLOADED_AUDIO_ENDPOINT = "http://34.246.239.57/process_uploaded_audio"
    DIARIZED_SUMMARY_ENDPOINT = "http://34.246.239.57/diarized_summary"

    @st.cache_resource(show_spinner=False)
    def predict(audio_file):
        files = {'input_file': (audio_file.name, audio_file, 'application/octet-stream')}
        response = requests.post(PROCESS_UPLOADED_AUDIO_ENDPOINT, files=files)
        response = response.json()
        print(response)
        return response
    
    @st.cache_resource(show_spinner=False)
    def diarized_summary(diarized_transcript):
        response = requests.post(DIARIZED_SUMMARY_ENDPOINT, json={"transcript": diarized_transcript},)
        response = response.json()
        print(response)
        return response
    
    def create_pdf(output,num):

        # Create a new PDF file in memory
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)

        # Set the starting y-coordinate
        y_position = 750

        # Define the frame and paragraph style
        frame_width = 500
        frame_height = 20
        styles = getSampleStyleSheet()
        paragraph_style = styles['Normal']
        paragraph_style.fontSize = 10
        paragraph_style.textColor = colors.black

        if num == 1:
            download_label = "Transcript"
            for sentence in output:
                p = Paragraph(sentence, paragraph_style)
                wrapped_width, wrapped_height = p.wrap(frame_width, frame_height)
                
                # Add extra space for truncated sentences
                if wrapped_height > frame_height:
                    extra_space = 5  # Adjust this value to change the extra space
                    y_position -= extra_space

                p.drawOn(c, 50, y_position)
                y_position -= 20  # Adjust this value to change the spacing between sentences

                # Add extra space for truncated sentences
                if wrapped_height > frame_height:
                    y_position -= extra_space

                if y_position < 20:  # If the next line will be off the page, create a new page
                    c.showPage()
                    y_position = 750

        elif num == 2:
            download_label = "Summary"
            styles = getSampleStyleSheet()
            paragraph_style = styles['Normal']
            paragraph_style.fontSize = 12
            paragraph_style.leading = 14
            paragraph_style.alignment = TA_LEFT
            paragraph = Paragraph(output, paragraph_style)
            paragraph.wrapOn(c, 500, 800)
            paragraph.drawOn(c, 40, 600)

        elif num == 3:
            download_label = "Diarized Transcript"
            # Write each sentence to the PDF
            for sentence in output:
                p = Paragraph(sentence, paragraph_style)
                wrapped_width, wrapped_height = p.wrap(frame_width, frame_height)
                
                # Add extra space for truncated sentences
                if wrapped_height > frame_height:
                    extra_space = 5  # Adjust this value to change the extra space
                    y_position -= extra_space

                p.drawOn(c, 50, y_position)
                y_position -= 20  # Adjust this value to change the spacing between sentences

                # Add extra space for truncated sentences
                if wrapped_height > frame_height:
                    y_position -= extra_space

                if y_position < 20:  # If the next line will be off the page, create a new page
                    c.showPage()
                    y_position = 750

        elif num == 4:
            download_label = "Diarized Summary"
            styles = getSampleStyleSheet()
            paragraph_style = styles['Normal']
            paragraph_style.fontSize = 12
            paragraph_style.leading = 14
            paragraph_style.alignment = TA_LEFT
            paragraph = Paragraph(output, paragraph_style)
            paragraph.wrapOn(c, 500, 800)
            paragraph.drawOn(c, 40, 600)



        c.showPage()
        # Save and close the PDF
        c.save()

        pdf_buffer.seek(0)

        st.download_button(
            label=download_label,
            data=pdf_buffer,
            file_name=download_label+".pdf",
            mime="application/pdf",
        )
    
    def create_word_doc(output, num):

        # Create a new Word document
        doc = Document()

        if num == 1:
            download_label = "Transcript"
            doc.add_paragraph(output)
        elif num == 2:
            download_label = "Summary"
            doc.add_paragraph(output)
        elif num == 3:
            download_label = "Diarized Transcript"
            # Write each sentence in the output_list as a separate paragraph
            for sentence in output:
                doc.add_paragraph(sentence)
        elif num == 4:
            download_label = "Diarized Summary"
            doc.add_paragraph(output)


        # Save the document to a BytesIO object
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        st.download_button(
            label=download_label,
            data=doc_buffer,
            file_name=download_label+".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def create_text_file(content, num):

        if num == 1:
            download_label = "Transcript"
        elif num == 2:
            download_label = "Summary"
        elif num == 3:
            download_label = "Diarized Transcript"
            # Convert the list of sentences to a single string with newlines between each sentence
            content = "\n".join(content)
        elif num == 4:
            download_label = "Diarized Summary"

        # Send the file to the user to download
        st.download_button(
            label=download_label,
            data=content,
            file_name=download_label+".txt",
            mime="text/plain",
        )

    def save_to_bucket(user_identity, content, filename):

        try:
            s3_key = f"{user_identity}/{filename}"
            s3_client.put_object(
                Bucket=S3_BUCKET,
                Key=s3_key,
                Body=content,
                ContentType="text/plain"
            )

            print(f"File {filename} uploaded to S3 bucket {S3_BUCKET}.")

        except Exception as e:

            print(f"Error saving to S3: {e}")

            raise



    if audio_file is not None:

        st.audio(audio_file)
        
        with st.spinner("Transcribing Audio..."):
            file_title = audio_file.name.rsplit('.', 1)[0]
            response = predict(audio_file)
            transcript = response["transcript"]
            diarization = response["diarization"]
            summary = response["summary"]


        st.markdown("### Transcription")
        with st.expander("**Transcript**", expanded=True):
            st.write(transcript)


        st.markdown("### Summarization")
        with st.expander("**Summary**", expanded=True):
                st.write(summary[0]["Summary"])


        st.markdown("### Diarization")
        # Get the speaker tag and custom name from the user
        speaker_tag = st.text_input("Speaker Label To Replace:", "")
        custom_name = st.text_input("Replacement Name:", "")
        # Add a button to trigger the replacement
        replace_button = st.button("Replace")
        st.write("")
        saved_diarized_sentences = []
        conversation_str = ""  


        with st.expander("**Speaker Identification**", expanded=True):
            if "sentencesthree" not in st.session_state:

                st.session_state.sentencesthree = []
                current_sentence = ''

                for character in diarization:
                    current_sentence += character
                    if character in ['.', '?', '!']:
                        st.session_state.sentencesthree.append(current_sentence.strip())
                        current_sentence = ''

            # If the button is clicked and both inputs are not empty, replace the speaker tag with the custom name
            if replace_button and speaker_tag and custom_name:
                st.session_state.sentencesthree = [s.replace(speaker_tag, custom_name) for s in st.session_state.sentencesthree]


            # Display the text with custom speaker names and each sentence on a new line
            for sentence in st.session_state.sentencesthree:
                st.write(sentence)
                saved_diarized_sentences.append(sentence)
                conversation_str += sentence

        
        st.write("")

        col1, col2 = st.columns([3.5,6])
        colA, colB = st.columns([12,10])

        with col1:
            
            st.write("")

            with colA:
                st.write("")

        with col2:

            st.markdown("**_Summarize transcript while including speaker(s) identity_**")
            
            with colB:
                expand = st.button("**Summarize**")

        
        with st.expander("**Diarized Summary**",expanded=True):
            if expand:    
                with st.spinner("Summarizing Diarized Transcription..."):
                    response = diarized_summary(conversation_str)
                    summary_diarized = response["Summary"]
                    st.session_state.diarize_summarythree = summary_diarized[0]["Summary"]
                    st.write(st.session_state.diarize_summarythree)
        

        st.markdown("### Save Transcription")
        with st.expander("**Save to File/Database**",expanded=True):

            file_type = st.selectbox("Select Format", ("DOCX", "PDF", "TXT", "DATABASE"), index=0)
            download_button = st.button("Confirm")

            if download_button:
                if file_type == "PDF":
                    st.divider()
                    st.write("**Files Available for Download**")
                    sentTra = re.split('(?<=[.!?]) +', transcript)
                    create_pdf(sentTra,1)
                    create_pdf(summary[0]["Summary"],2)
                    create_pdf(saved_diarized_sentences,3)
                    if 'diarize_summarythree' in st.session_state:  
                        create_pdf(st.session_state.diarize_summarythree, 4)

                elif file_type == "DOCX":
                    st.divider()
                    st.write("**Files Available for Download**")
                    create_word_doc(transcript, 1)
                    create_word_doc(summary[0]["Summary"], 2)
                    create_word_doc(saved_diarized_sentences, 3)
                    if 'diarize_summarythree' in st.session_state:  
                        create_word_doc(st.session_state.diarize_summarythree, 4)

                elif file_type == "TXT":
                    st.divider()
                    st.write("**Files Available for Download**")
                    create_text_file(transcript, 1)
                    create_text_file(summary[0]["Summary"], 2)
                    create_text_file(saved_diarized_sentences, 3)
                    if 'diarize_summarythree' in st.session_state:  
                        create_text_file(st.session_state.diarize_summarythree, 4)

                elif file_type == "DATABASE":
                    st.write("**_Successfully Saved To Database_**")
                    save_to_bucket(user_identity, transcript, "Transcript - "+file_title)
                    save_to_bucket(user_identity, summary[0]["Summary"], "Summary - "+file_title)
                    content = "\n".join(saved_diarized_sentences).encode("utf-8")
                    save_to_bucket(user_identity, content, "Diarized Transcript - "+file_title)
                    if 'diarize_summarythree' in st.session_state:  
                        save_to_bucket(user_identity, st.session_state.diarize_summarythree, "Diarized Summary - "+file_title)

else:
        
    st.write("**Login For Access**")
